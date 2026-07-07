"""
ml/parser.py
Resume parsing pipeline.

Steps:
1. Extract raw text from PDF or DOCX
2. Extract skills using spaCy NER + pattern matching
3. Normalise skills (ReactJS -> React)
4. Extract years of experience from date ranges
5. Classify visa status from keywords
6. Return structured dict ready to store in candidates table

Usage:
    from ml.parser import parse_resume

    with open("resume.pdf", "rb") as f:
        result = parse_resume(f.read(), "pdf")
    # result = {name, email, phone, skills, yoe, visa_status, raw_text}
"""
import io
import re
import logging
from pathlib import Path
from typing import Optional

import spacy

logger = logging.getLogger(__name__)

# ── Skill normalisation map ───────────────────────────────────────────────────
SKILL_ALIASES = {
    # JavaScript ecosystem
    "reactjs": "React", "react.js": "React", "react js": "React",
    "nodejs": "Node.js", "node js": "Node.js", "node": "Node.js",
    "vuejs": "Vue.js", "vue js": "Vue.js",
    "angularjs": "Angular", "angular js": "Angular",
    "nextjs": "Next.js", "next js": "Next.js",
    "typescript": "TypeScript", "ts": "TypeScript",
    "javascript": "JavaScript", "js": "JavaScript",

    # Python
    "python3": "Python", "python 3": "Python",
    "fastapi": "FastAPI", "fast api": "FastAPI",
    "django rest framework": "Django", "drf": "Django",
    "flask": "Flask",

    # ML/AI
    "machine learning": "Machine Learning", "ml": "Machine Learning",
    "deep learning": "Deep Learning", "dl": "Deep Learning",
    "natural language processing": "NLP", "nlp": "NLP",
    "computer vision": "Computer Vision", "cv": "Computer Vision",
    "scikit learn": "scikit-learn", "sklearn": "scikit-learn",
    "tensorflow": "TensorFlow", "tf": "TensorFlow",
    "pytorch": "PyTorch", "torch": "PyTorch",

    # Cloud
    "amazon web services": "AWS", "aws": "AWS",
    "google cloud platform": "GCP", "gcp": "GCP", "google cloud": "GCP",
    "microsoft azure": "Azure", "azure": "Azure",

    # Data
    "postgresql": "PostgreSQL", "postgres": "PostgreSQL",
    "mysql": "MySQL", "ms sql": "SQL Server", "mssql": "SQL Server",
    "mongodb": "MongoDB", "mongo": "MongoDB",
    "apache spark": "Spark", "pyspark": "Spark",
    "apache kafka": "Kafka",
    "apache airflow": "Airflow",
    "elasticsearch": "Elasticsearch", "elastic search": "Elasticsearch",

    # DevOps
    "kubernetes": "Kubernetes", "k8s": "Kubernetes",
    "docker": "Docker",
    "terraform": "Terraform",
    "ci/cd": "CI/CD", "cicd": "CI/CD",
    "jenkins": "Jenkins",
    "github actions": "GitHub Actions",

    # Java
    "java spring": "Spring Boot", "spring framework": "Spring Boot",
    "spring boot": "Spring Boot",

    # Other
    "c sharp": "C#", "c#": "C#", ".net": ".NET", "dotnet": ".NET",
    "golang": "Go", "go lang": "Go",
    "rust lang": "Rust",
    "rest api": "REST API", "restful": "REST API",
    "graphql": "GraphQL",
    "snowflake": "Snowflake",
    "dbt": "dbt",
    "power bi": "Power BI", "powerbi": "Power BI",
    "tableau": "Tableau",
    "salesforce": "Salesforce", "sfdc": "Salesforce",
    "sap": "SAP",
}

# Known tech skills for pattern matching (lowercase)
KNOWN_SKILLS = {
    "python", "java", "javascript", "typescript", "react", "node.js", "angular",
    "vue.js", "next.js", "aws", "azure", "gcp", "docker", "kubernetes", "terraform",
    "postgresql", "mysql", "mongodb", "redis", "kafka", "spark", "airflow",
    "machine learning", "deep learning", "nlp", "computer vision", "scikit-learn",
    "tensorflow", "pytorch", "xgboost", "lightgbm", "pandas", "numpy",
    "fastapi", "django", "flask", "spring boot", "go", "rust", "c#", ".net",
    "scala", "kotlin", "swift", "sql", "nosql", "graphql", "rest api",
    "snowflake", "dbt", "power bi", "tableau", "salesforce", "sap",
    "ci/cd", "jenkins", "github actions", "git", "linux", "bash",
    "elasticsearch", "dynamodb", "cassandra", "oracle", "hive", "hadoop",
}

# Visa status keywords
VISA_PATTERNS = {
    "citizen":  ["us citizen", "u.s. citizen", "united states citizen", "american citizen", "usc"],
    "gc":       ["green card", "permanent resident", "gc holder", "lawful permanent"],
    "h1b":      ["h1b", "h-1b", "h1-b", "h 1b"],
    "opt":      ["opt", "f1 opt", "f-1 opt", "optional practical training"],
    "stem_opt": ["stem opt", "stem extension", "stem opt extension"],
    "ead":      ["ead", "employment authorization", "employment auth"],
}


# ── Text extraction ───────────────────────────────────────────────────────────

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract plain text from PDF bytes using pdfminer.six."""
    from pdfminer.high_level import extract_text_to_fp
    from pdfminer.layout import LAParams

    output = io.StringIO()
    with io.BytesIO(file_bytes) as pdf_file:
        extract_text_to_fp(pdf_file, output, laparams=LAParams(), output_type="text", codec="utf-8")
    return output.getvalue()


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract plain text from DOCX bytes using python-docx."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)


def extract_text(file_bytes: bytes, file_type: str) -> str:
    """Route to the correct extractor based on file type."""
    ft = file_type.lower().lstrip(".")
    if ft == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ft in ("docx", "doc"):
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {file_type}. Use 'pdf' or 'docx'.")


# ── Contact extraction ────────────────────────────────────────────────────────

def extract_email(text: str) -> Optional[str]:
    match = re.search(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}", text)
    return match.group(0).lower() if match else None


def extract_phone(text: str) -> Optional[str]:
    match = re.search(
        r"(\+?1?\s?)?(\(?\d{3}\)?[\s.\-]?)(\d{3}[\s.\-]?\d{4})", text
    )
    return match.group(0).strip() if match else None


def extract_name_spacy(text: str, nlp) -> Optional[str]:
    """Use spaCy PERSON entity to guess the candidate's name (first PERSON entity)."""
    doc = nlp(text[:500])  # name is almost always near the top
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            return ent.text.strip()
    # Fallback: first non-empty line that looks like a name
    for line in text.split("\n")[:5]:
        line = line.strip()
        if 2 <= len(line.split()) <= 4 and line.replace(" ", "").isalpha():
            return line
    return None


# ── Skill extraction ──────────────────────────────────────────────────────────

def normalise_skill(raw: str) -> str:
    key = raw.lower().strip()
    return SKILL_ALIASES.get(key, raw.title())


def extract_skills(text: str, nlp) -> list[str]:
    """
    Two-pass skill extraction:
    1. Pattern match against KNOWN_SKILLS
    2. spaCy NER for any PRODUCT/ORG entities that look like tech
    """
    text_lower = text.lower()
    found = set()

    # Pass 1: direct keyword match
    for skill in KNOWN_SKILLS:
        # Use word boundary matching
        pattern = r"\b" + re.escape(skill) + r"\b"
        if re.search(pattern, text_lower):
            found.add(normalise_skill(skill))

    # Pass 2: check aliases
    for alias in SKILL_ALIASES:
        pattern = r"\b" + re.escape(alias) + r"\b"
        if re.search(pattern, text_lower):
            found.add(SKILL_ALIASES[alias])

    return sorted(found)


# ── Years of experience ───────────────────────────────────────────────────────

def extract_yoe(text: str) -> float:
    """
    Estimate total years of experience from date ranges in resume text.
    Looks for patterns like "Jan 2018 – Mar 2022" or "2019 - 2023".
    """
    import datetime

    # Match year ranges: 2018 - 2022, 2018 – Present, etc.
    year_range_pattern = re.compile(
        r"((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+)?(\d{4})\s*[-–—]\s*((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+)?(\d{4}|present|current|now)",
        re.IGNORECASE,
    )

    current_year = datetime.datetime.now().year
    total_months = 0
    seen_ranges = []

    for match in year_range_pattern.finditer(text):
        start_year = int(match.group(2))
        end_str = match.group(4).lower()
        end_year = current_year if end_str in ("present", "current", "now") else int(end_str)

        if 1980 <= start_year <= current_year and start_year <= end_year <= current_year + 1:
            # Avoid double-counting overlapping ranges
            overlap = False
            for s, e in seen_ranges:
                if not (end_year <= s or start_year >= e):
                    overlap = True
                    break
            if not overlap:
                total_months += (end_year - start_year) * 12
                seen_ranges.append((start_year, end_year))

    if total_months > 0:
        return round(total_months / 12, 1)

    # Fallback: look for explicit "X years of experience" statements
    explicit = re.search(r"(\d+)\+?\s+years?\s+(?:of\s+)?experience", text, re.IGNORECASE)
    if explicit:
        return float(explicit.group(1))

    return 0.0


# ── Visa status ───────────────────────────────────────────────────────────────

def extract_visa_status(text: str) -> str:
    """
    Returns one of: citizen, gc, h1b, opt, stem_opt, ead, unknown.
    Checks STEM OPT before OPT to avoid false matches.
    """
    text_lower = text.lower()
    for status in ["stem_opt", "citizen", "gc", "h1b", "ead", "opt"]:
        for keyword in VISA_PATTERNS[status]:
            if keyword in text_lower:
                return status
    return "unknown"


# ── Main parser ───────────────────────────────────────────────────────────────

_nlp = None  # lazy-loaded spaCy model


def _get_nlp():
    global _nlp
    if _nlp is None:
        try:
            _nlp = spacy.load("en_core_web_sm")
            logger.info("Loaded spaCy model: en_core_web_sm")
        except OSError:
            logger.warning("spaCy model not found. Run: python -m spacy download en_core_web_sm")
            _nlp = spacy.blank("en")
    return _nlp


def parse_resume(file_bytes: bytes, file_type: str) -> dict:
    """
    Parse a resume from raw bytes.

    Args:
        file_bytes: raw PDF or DOCX file bytes
        file_type:  "pdf" or "docx"

    Returns dict with keys:
        name, email, phone, skills, yoe, visa_status, raw_text
    """
    nlp = _get_nlp()

    # 1. Extract raw text
    try:
        raw_text = extract_text(file_bytes, file_type)
    except Exception as e:
        logger.error(f"Text extraction failed: {e}")
        raise

    if not raw_text.strip():
        raise ValueError("Could not extract any text from the resume. Check if it's a scanned image PDF.")

    # 2. Extract all fields
    name       = extract_name_spacy(raw_text, nlp)
    email      = extract_email(raw_text)
    phone      = extract_phone(raw_text)
    skills     = extract_skills(raw_text, nlp)
    yoe        = extract_yoe(raw_text)
    visa_status = extract_visa_status(raw_text)

    result = {
        "name":        name or "Unknown",
        "email":       email,
        "phone":       phone,
        "skills":      skills,
        "yoe":         yoe,
        "visa_status": visa_status,
        "raw_text":    raw_text,
    }

    logger.info(f"Parsed resume: {name} | skills={len(skills)} | yoe={yoe} | visa={visa_status}")
    return result


def parse_resume_file(file_path: str) -> dict:
    """Convenience wrapper — parse directly from a file path."""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    with open(path, "rb") as f:
        return parse_resume(f.read(), path.suffix.lstrip("."))
